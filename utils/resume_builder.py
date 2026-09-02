"""Build downloadable .docx documents for the tailoring workflow.

Creates a tailored resume document (with professional summary, key skills,
experience highlights, the full original resume for reference, and any
honesty notes) as well as a separate cover letter document. These files are
what get served to the end user after a resume is tailored to a job.
"""

import logging
from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


class ResumeBuilder:
    """Generates a downloadable .docx tailored resume + cover letter."""

    @staticmethod
    def build_tailored_resume(
        candidate_name: str,
        original_resume_text: str,
        tailored: Dict,
        output_path: Path,
    ) -> Path:
        doc = Document()

        title = doc.add_heading(candidate_name or "Candidate Resume", level=0)

        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(tailored.get("professional_summary", ""))

        skills = tailored.get("highlighted_skills") or []
        if skills:
            doc.add_heading("Key Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        bullets = tailored.get("tailored_experience_bullets") or []
        if bullets:
            doc.add_heading("Experience Highlights", level=1)
            for bullet in bullets:
                doc.add_paragraph(bullet, style="List Bullet")

        doc.add_heading("Full Original Resume Text (reference)", level=1)
        p = doc.add_paragraph(original_resume_text)
        for run in p.runs:
            run.font.size = Pt(9)

        honesty_notes = tailored.get("honesty_notes")
        if honesty_notes:
            doc.add_heading("⚠ Please verify before submitting", level=1)
            doc.add_paragraph(honesty_notes)

        doc.save(str(output_path))
        return output_path

    @staticmethod
    def build_cover_letter(cover_letter_text: str, output_path: Path) -> Path:
        doc = Document()
        doc.add_heading("Cover Letter", level=0)
        for paragraph in cover_letter_text.split("\n\n"):
            doc.add_paragraph(paragraph.strip())
        doc.save(str(output_path))
        return output_path
